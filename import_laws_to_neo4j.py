from docx import Document
from neo4j import GraphDatabase
import os
import re

# Neo4j Database settings
NEO4J_URI = "bolt://localhost:7687"
NEO4J_USER = "neo4j"
NEO4J_PASSWORD = "test1234"

# Folder containing .docx files
DOCX_FOLDER = "./laws"

# Connect to Neo4j
driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))

def clear_database(tx):
    tx.run("MATCH (n) DETACH DELETE n")

def create_law_node(tx, section_number, text, book, year, source):
    tx.run(
        """
        MERGE (l:Law {number: $number, book: $book, year: $year, source: $source})
        SET l.text = $text
        """,
        number=section_number, text=text, book=book, year=year, source=source
    )

def create_parent_child_relation(tx, parent_number, child_number):
    tx.run(
        """
        MATCH (parent:Law {number: $parent_number})
        MATCH (child:Law {number: $child_number})
        MERGE (parent)-[:HAS_CHILD]->(child)
        """,
        parent_number=parent_number, child_number=child_number
    )

def create_sibling_relation(tx, first_number, second_number):
    tx.run(
        """
        MATCH (a:Law {number: $first_number})
        MATCH (b:Law {number: $second_number})
        MERGE (a)-[:NEXT_SIBLING]->(b)
        """,
        first_number=first_number, second_number=second_number
    )

def create_refers_to_relation(tx, from_number, to_number):
    tx.run(
        """
        MATCH (a:Law {number: $from_number})
        MATCH (b:Law {number: $to_number})
        MERGE (a)-[:REFERS_TO]->(b)
        """,
        from_number=from_number, to_number=to_number
    )

# def parse_docx(filepath, filename):
#     document = Document(filepath)
#     lines = []
#     for para in document.paragraphs:
#         text = para.text.strip()
#         if text:
#             match = re.match(r'^(\d+(\.\d+)*)\s+(.*)', text)
#             if match:
#                 section_number = match.group(1)
#                 section_text = match.group(3)
#                 lines.append((section_number, section_text))
#     return lines

def parse_docx(filepath, filename):
    document = Document(filepath)
    lines = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            # Match things like '1.', '1.1.', '1.1.1.', followed by text
            match = re.match(r'^(\d+(\.\d+)*\.)\s+(.*)', text)
            if match:
                section_number = match.group(1).rstrip(".")  # remove trailing dot
                section_text = match.group(3)
                lines.append((section_number, section_text))
    return lines

def extract_metadata_from_filename(filename):
    parts = filename.replace(".docx", "").split("_")
    book = parts[0]
    year = parts[1]
    source = parts[2]
    return book, year, source

def main():
    with driver.session(database="neo4j") as session:
        print("✅ Connected to Neo4j")

        # Step 1: Clear database
        session.execute_write(clear_database)

        for filename in os.listdir(DOCX_FOLDER):
            if filename.endswith(".docx") or filename.endswith(".doc"):
                filepath = os.path.join(DOCX_FOLDER, filename)
                print(f"Processing {filename}")

                book, year, source = extract_metadata_from_filename(filename)
                sections = parse_docx(filepath, filename)
                print(f"Found {len(sections)} sections in {filename}")

                last_section_per_level = {}  # for siblings

                for idx, (section_number, section_text) in enumerate(sections):
                    session.execute_write(create_law_node, section_number, section_text, book, year, source)

                    # Handle Parent-Child
                    parent_number = ".".join(section_number.split(".")[:-1])
                    if parent_number:
                        session.execute_write(create_parent_child_relation, parent_number, section_number)

                    # Handle Sibling
                    level = section_number.count(".")
                    if level in last_section_per_level:
                        previous_sibling = last_section_per_level[level]
                        session.execute_write(create_sibling_relation, previous_sibling, section_number)

                    last_section_per_level[level] = section_number

                # Step 2: Handle internal references (See Section X.X)
                for section_number, section_text in sections:
                    references = re.findall(r'See Section (\d+(\.\d+)*)', section_text)
                    for ref_number, _ in references:
                        session.execute_write(create_refers_to_relation, section_number, ref_number)

    print("✅ Import complete!")

if __name__ == "__main__":
    main()